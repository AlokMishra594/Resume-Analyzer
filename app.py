# THIS IS A RESUME ANALYZER, IT PREDICTS THE JOB ROLE ACCORDING TO YOUR RESUME.

# import numpy as np
# import pandas as pd
# import spacy

import os
port = int(os.environ.get("PORT",8501))
import re
import nltk
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.model_selection import train_test_split
from sklearn.linear_model import LogisticRegression
from sklearn.metrics import accuracy_score
import pickle

import pdfminer
from pdfminer.high_level import extract_text
import docx

# you can use this in colab to download the dataset.
# !kaggle datasets download -d snehaanbhawal/resume-dataset -p /content/ --unzip

# df = pd.read_csv('Resume.csv')
# df.head()

# nltk.download('stopwords')
from nltk.corpus import stopwords

try:
    stop_words = set(stopwords.words('english'))
except:
    nltk.download('stopwords')
    stop_words = set(stopwords.words('english'))
# stop_words = set(stopwords.words('english'))

model = pickle.load(open("model.pkl","rb"))
tfidf = pickle.load(open("vectorizer.pkl","rb"))

def clean_text(text):
    text = text.lower()
    text = re.sub(r'\W', ' ', text)  # remove special characters
    text = re.sub(r'\d', '', text)   # remove numbers
    words = text.split()
    words = [word for word in words if word not in stop_words]
    return " ".join(words)

# df['cleaned_resume'] = df['Resume_str'].apply(clean_text)
# tfidf = TfidfVectorizer(max_features=5000)
# X = tfidf.fit_transform(df['cleaned_resume']).toarray()

# y = df['Category']
# X_train, X_test, y_train, y_test = train_test_split(X, y, test_size=0.2, random_state=42)

# model = LogisticRegression()
# model.fit(X_train, y_train)

# y_pred = model.predict(X_test)
# # print("Accuracy:", accuracy_score(y_test, y_pred))

# pickle.dump(model, open('model.pkl', 'wb'))
# pickle.dump(tfidf, open('vectorizer.pkl', 'wb'))


# def extract_text_from_pdf(file):
#     return extract_text(file)

# def extract_text_from_docx(file):
#     doc = docx.Document(file)
#     return " ".join([para.text for para in doc.paragraphs])

# def predict_resume(file_path):
#     if file_path.type('.pdf'):
#         text = extract_text_from_pdf(file_path)
#     else:
#         text = extract_text_from_docx(file_path)

#     cleaned = clean_text(text)

#     vector = tfidf.transform([cleaned]).toarray()
#     prediction = model.predict(vector)

#     return prediction[0]

def predict_resume(uploaded_file):
    uploaded_file.seek(0)  # important

    if uploaded_file.type == "application/pdf":
        text = extract_text(uploaded_file)
    else:
        doc = docx.Document(uploaded_file)
        text = " ".join([para.text for para in doc.paragraphs])

    cleaned = clean_text(text)

    vector = tfidf.transform([cleaned]).toarray()
    prediction = model.predict(vector)

    return prediction[0]

import streamlit as st

st.title("Resume Analyzer")
st.text("Which Role is Best According to Your Resume")

uploaded_file = st.file_uploader("Upload Resume", type=['pdf', 'docx'])

if uploaded_file:
    # with open(uploaded_file.name, "wb") as f:
    #     f.write(uploaded_file.getbuffer())

    # result = predict_resume(uploaded_file.name)
    result = predict_resume(uploaded_file)

    st.success(f"Predicted Job Role: {result}")
